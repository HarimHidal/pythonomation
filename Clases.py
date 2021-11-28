# -*- coding: utf-8 -*-
"""
Created on Fri Jan 15 17:11:07 2021

@author: harim
"""

import webbrowser as wb
from os import path, startfile, listdir
from docx import Document
from time import ctime, sleep



Master_file = r"C:\Users\X\X\X\Master.docx"

log = sorted([ x for x in range(1,50) ], reverse=True)


# CLASSES #####################################################################


class Clase ():
    
    Horarios_docx = Master_file.replace("Master.docx", "Horarios.docx")
    Zoom_exe = r"C:\Users\X\X\X\X\X\Zoom.exe"
        
    def OpenClass (self, directory, parcial, link="", Zoom=True):
        """Opens the schedule and the .pptx located in the class directory; link for Zoom is optional."""
        
        print("[{}]:\tOPEN THE SCHOOL SCHEDULE DOCX...\n".format(log.pop()))
        startfile(self.Horarios_docx)
        
        nombre_clase = path.basename(directory)
        
        print("[{}]:\tOPEN THE DIRECTORY AND PPTX FOR '{}'...\n".format(log.pop(), nombre_clase))        
        ls = listdir(directory)
        for item in ls:    
            if item[0] == parcial: # and path.splitext(item)[1] == ".pptx":      
                startfile(path.join(directory, item))
                break
        startfile(directory)
        
        if Zoom == True:
            
            print("[{}]:\tOPEN ZOOM FOR THE CLASS...\n".format(log.pop()))
            startfile(self.Zoom_exe)
            
            print("[{}]:\tOPEN THE LINK FOR '{}'...\n".format(log.pop(), nombre_clase))
            wb.open(link)   


class ParseDocument ():
    
    filename = Master_file
    
    def ParseColumns (self, table_number):
        """Parses the .docx file to return a dictionary of lists with the values of the columns."""
        
        doc = Document(self.filename)
        table = doc.tables[table_number]
        
        result = {}
        for i, column in enumerate(table.columns):
            for j, cell in enumerate(column.cells):
                if j == 0:
                    result[cell.text] = result.get(cell.text, [])
                    header = cell.text
                else:
                    result[header].append(cell.text.strip())
                    
        return result
    
    
    def ParseRows (self, table_number):
        """Parses the .docx file to return a dictionary of lists with the values of the rows."""
        
        doc = Document(self.filename)
        table = doc.tables[table_number]
        
        result = {}
        for i, row in enumerate(table.rows):
            if i == 0:
                continue
            for j, cell in enumerate(row.cells):
                if j == 0:
                    result[cell.text] = result.get(cell.text, [])
                    header = cell.text
                else:    
                    result[header].append(cell.text.strip())
            
        return result    
        
        
# FUNCTIONS ###################################################################


def TraducirDia (dia):
    
    conversiones = {"Mon":"Lunes", "Tue":"Martes", "Wed":"Miercoles", "Thu":"Jueves", "Fri":"Viernes", "Sat":"Sabado", "Sun":"Domingo"}
    result = conversiones[dia]
    return result

        
print()########################################################################


# START OF DEFAULT EXECUTION ##################################################


dia, hora = TraducirDia(ctime()[:3]).upper(), int(ctime()[11:13])

reader = ParseDocument()
starter = Clase()

horario = reader.ParseColumns(0)
links = reader.ParseRows(1)
update = reader.ParseRows(2)


parcial = update["PARCIAL"][0]

base_directory = update["BASE"][0]


if dia in horario.keys():
    
    flag = False
    for i, elemento in enumerate(horario["HORARIO"]):
        
        if (int(elemento[0:2])-1) == hora:
            
            indice_clase = i
            flag = True
            break
    
    if flag == True:
        
        nombre_clase = horario[dia][indice_clase]
    
        if nombre_clase != "" and nombre_clase in links.keys():
            
            link_clase = links[nombre_clase][1]
    
            starter.OpenClass(path.join(base_directory, nombre_clase), parcial, link_clase)    
            
            if nombre_clase in update.keys():
                
                startfile(update[nombre_clase][0])

        else:        
            
            print("[{}]:\tYOU HAVE NO CLASS SCHEDULED...\n".format(log.pop()))
            
    else:
        
        print("[{}]:\tYOU HAVE NO CLASS SCHEDULED...\n".format(log.pop()))
        
else:

    print("[{}]:\tYOU HAVE NO CLASS SCHEDULED...\n".format(log.pop()))
    

# START OF USER INPUT EXECUTION ##############################################


opciones = {}
for index, element in enumerate(sorted(links.keys())):
    opciones[index+1] = element

print("************************************************************")
print("\nVALID COMMAND-LINE ARGUMENTS:\n")
print("    [ email ] --> abre el correo en el buscador")
print("    [ moodle ] --> abre Moodle en el buscador")
print("    [ clase ] --> forzar la clase especificada\n")
print("PRESS 'Enter' TO EXIT...")


while True:  

    command_line = input("\n>> ")
    
    if command_line == "":
        break
    
    elif command_line.isalpha():
       
        if command_line in "email":
            wb.open("https://mail.google.com/mail/u/XXXXXX@XX.XXX.XX/#inbox")
            break
            
        elif command_line in "moodle":   
            wb.open("https://moodle.XX.XXX.XX/my/")
            break
        
        elif command_line in "clase":      
            print()
            for opcion in opciones:
                print("    {} --> {}".format(opcion, opciones[opcion]))
            continue
            
        elif command_line in [ x.lower() for x in opciones.values() ]:

            for key, value in opciones.items():
                if command_line == value.lower():
                    command_line = opciones[key]
                    break
            
            print()
            starter.OpenClass(path.join(base_directory, command_line), parcial, links[command_line][1], Zoom=False)    
            if command_line in update.keys():
                startfile(update[command_line][0])
            break
    
        else: 
            print("\nERROR: Enter a valid command.")
    
    elif command_line.isnumeric():
       
        try:
            command_line = opciones[int(command_line)]
        except:
            print("\nERROR: Enter a valid class number.")
            continue
       
        if command_line in links.keys():
            
            print()
            starter.OpenClass(path.join(base_directory, command_line), parcial, links[command_line][1], Zoom=False)      
            if command_line in update.keys():
                startfile(update[command_line][0])
            break
    
    else:
        print("\nERROR: Enter a valid command.")
        
sleep(1)

        
        